import streamlit as st
import pdfplumber
import docx
import tempfile
import os
import google.generativeai as genai

# ---------------- DARK THEME STYLING ----------------
st.markdown('''
    <style>
        .stApp {
            background: linear-gradient(135deg, #121212 0%, #1e1e1e 100%);
            color: #e0e0e0;
        }
        .main { background-color: transparent; }
        .stTitle {
            color: #ff79c6; font-size: 2.8rem; font-weight: 800;
            text-align: center; margin-bottom: 1rem; animation: fadeIn 1s ease-in-out;
        }
        .step-card {
            background: rgba(30, 30, 30, 0.9); padding: 1.5rem;
            border-radius: 15px; box-shadow: 0 4px 20px rgba(255, 121, 198, 0.15);
            margin-bottom: 1.5rem; animation: fadeInUp 0.8s ease-in-out;
        }
        .stHeader { color: #bd93f9; font-size: 1.4rem; font-weight: 600; margin-bottom: 0.8rem; }
        .stTextArea textarea {
            background: #1e1e1e; border-radius: 10px; border: 1px solid #444;
            font-size: 1.05rem; padding: 0.7rem; color: #f8f8f2;
        }
        .stTextArea textarea:focus {
            border-color: #ff79c6; box-shadow: 0 0 6px rgba(255, 121, 198, 0.5);
        }
        .stButton button, .stDownloadButton button {
            font-weight: 600; border-radius: 8px; padding: 0.6em 1.5em;
            font-size: 1.1rem; border: none; transition: all 0.3s ease;
        }
        .stButton button { background-color: #ff79c6; color: #fff; }
        .stButton button:hover { background-color: #e066b4; transform: translateY(-2px); }
        .stDownloadButton button { background-color: #bd93f9; color: #fff; }
        .stDownloadButton button:hover { background-color: #a67ce6; transform: translateY(-2px); }
        ::-webkit-scrollbar { width: 8px; }
        ::-webkit-scrollbar-thumb { background: #ff79c6; border-radius: 10px; }
        @keyframes fadeIn { from { opacity: 0; transform: translateY(-10px); } to { opacity: 1; transform: translateY(0); } }
        @keyframes fadeInUp { from { opacity: 0; transform: translateY(20px); } to { opacity: 1; transform: translateY(0); } }
    </style>
''', unsafe_allow_html=True)

st.markdown('<h1 class="stTitle">Tailored Resume & Cover Letter Generator</h1>' \
            '<h3 class="stTitle">Craft your career narrative.</h3>', unsafe_allow_html=True)

# ---------------- API KEY ----------------
api_key = "AIzaSyAxQT-WD3hvNJvzn70zPddHPdfOXFmcn-Q"

# ---------------- SESSION STATE INIT ----------------
if "resume_docx" not in st.session_state:
    st.session_state.resume_docx = None
if "cover_docx" not in st.session_state:
    st.session_state.cover_docx = None

# ---------------- STEP 1 ----------------
st.markdown('<hr>', unsafe_allow_html=True)
st.markdown('<div class="stHeader">Step 1: Upload your resume (PDF/DOCX) or paste LinkedIn text</div>', unsafe_allow_html=True)

resume_file = st.file_uploader("Upload Your Resume", type=["pdf", "docx"])
resume_text = ""

if resume_file:
    if resume_file.name.lower().endswith(".pdf"):
        with pdfplumber.open(resume_file) as pdf:
            resume_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif resume_file.name.lower().endswith(".docx"):
        doc = docx.Document(resume_file)
        resume_text = "\n".join([para.text for para in doc.paragraphs])

st.text_area("Or paste your LinkedIn profile text here:", key="linkedin", value=resume_text, height=200)
st.markdown('</div>', unsafe_allow_html=True)

# ---------------- STEP 2 ----------------
st.markdown('<hr>', unsafe_allow_html=True)
st.markdown('<div class="stHeader">Step 2: Paste the job description</div>', unsafe_allow_html=True)
job_desc = st.text_area("Job Description", height=200)
st.markdown('</div>', unsafe_allow_html=True)

# ---------------- SAVE DOCX FUNCTION ----------------
def save_docx(text, filename):
    tmp_path = os.path.join(tempfile.gettempdir(), filename)
    doc = docx.Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(tmp_path)
    return tmp_path

# ---------------- STEP 3 ----------------
st.markdown('<hr>', unsafe_allow_html=True)
if st.button("🚀 Generate Tailored Resume & Cover Letter"):
    if not api_key:
        st.error("Please enter your Gemini API key in the sidebar.")
    elif not (resume_text or st.session_state.linkedin):
        st.error("Please upload a resume or paste LinkedIn text.")
    elif not job_desc:
        st.error("Please paste a job description.")
    else:
        # Remove old files if they exist
        if st.session_state.resume_docx and os.path.exists(st.session_state.resume_docx):
            os.remove(st.session_state.resume_docx)
        if st.session_state.cover_docx and os.path.exists(st.session_state.cover_docx):
            os.remove(st.session_state.cover_docx)

        st.info("⏳ Generating... This may take a moment.")
        resume_input = st.session_state.linkedin if st.session_state.linkedin else resume_text
        prompt = f"""
        You are an expert career consultant with 15+ years of experience in crafting ATS-friendly resumes and persuasive cover letters.
        TASK:
        Given the candidate’s resume and a specific job description, create:
        
        1. Tailored Resume:
           - Prioritize experiences, achievements, and skills directly relevant to the job description.
           - Use clear, professional, and concise language.
           - Optimize for Applicant Tracking Systems (ATS) with relevant keywords from the job description.
           - Maintain a clean and consistent structure with logical section ordering.
           - Quantify achievements wherever possible.
           - Do not invent skills, experiences, or qualifications not present in the original resume.
           - You may reword, reframe, and reorder content for maximum impact.
           - Avoid generic filler phrases.

        2. Matching Cover Letter:
           - Begin with an engaging, role-relevant hook that grabs attention in the first line. Avoid “Dear Hiring Manager”.
           - Address the company and role directly in the first paragraph.
           - Maintain a confident, energetic, and professional tone with short, active sentences.
           - Show, don’t just tell — use 1-2 concrete examples of achievements relevant to the job description.
           - End with a concise call-to-action inviting an interview and expressing excitement to contribute.
           - Keep it 3-4 paragraphs max, within one page.

        INPUTS:
        Candidate Resume: {resume_input}
        Job Description: {job_desc}

        OUTPUT FORMAT (exactly as below):
        ---RESUME---
        <tailored resume text>
        ---COVER LETTER---
        <tailored cover letter text>
        """
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-2.0-flash')
            response = model.generate_content(prompt)
            output = response.text if hasattr(response, 'text') else str(response)

            resume_out, cover_out = "", ""
            if "---RESUME---" in output and "---COVER LETTER---" in output:
                parts = output.split("---COVER LETTER---")
                resume_out = parts[0].replace("---RESUME---", "").strip()
                cover_out = parts[1].strip()
            else:
                resume_out = output

            st.subheader("📄 Tailored Resume")
            st.text_area("Resume", resume_out, height=300)
            st.subheader("💌 Cover Letter")
            st.text_area("Cover Letter", cover_out, height=300)
            st.balloons()

            # Save files and keep in session state
            st.session_state.resume_docx = save_docx(resume_out, "resume.docx")
            st.session_state.cover_docx = save_docx(cover_out, "cover_letter.docx")

        except Exception as e:
            st.error(f"Error: {e}")

# ---------------- DOWNLOAD BUTTONS ----------------
if st.session_state.resume_docx and st.session_state.cover_docx:
    with open(st.session_state.resume_docx, "rb") as f:
        st.download_button("⬇ Download Resume as DOCX", f, file_name="resume.docx")
    with open(st.session_state.cover_docx, "rb") as f:
        st.download_button("⬇ Download Cover Letter as DOCX", f, file_name="cover_letter.docx")

st.markdown('<hr>', unsafe_allow_html=True)
